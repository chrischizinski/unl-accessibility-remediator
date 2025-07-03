"""
Word Document (.docx) Accessibility Processor for UNL Accessibility Remediator

This module analyzes Word documents for WCAG 2.1 Level AA compliance issues
and provides recommendations for improvement.
"""

import os
import json
from pathlib import Path
from typing import Dict, List, Optional, Tuple, Any
import logging

# Word document processing
try:
    from docx import Document
    from docx.shared import RGBColor, Pt
    from docx.enum.style import WD_STYLE_TYPE
except ImportError as e:
    logging.warning(f"Word document processing libraries not available: {e}")
    Document = None

# Text analysis
try:
    import textstat
except ImportError:
    textstat = None

from .contrast_checker import ContrastChecker

logger = logging.getLogger(__name__)

class DocxAccessibilityProcessor:
    """Analyzes Word documents (.docx) for accessibility compliance"""
    
    def __init__(self):
        self.contrast_checker = ContrastChecker()
        self.issues = []
        self.fixes_applied = []
        
    def analyze_docx(self, file_path: str, apply_fixes: bool = False) -> Dict[str, Any]:
        """
        Analyze a Word document for accessibility issues
        
        Args:
            file_path: Path to the .docx file
            apply_fixes: Whether to attempt automatic fixes
            
        Returns:
            Dictionary containing analysis results
        """
        logger.info(f"Starting Word document accessibility analysis: {file_path}")
        
        if not Document:
            return self._create_error_result("python-docx library not available")
        
        try:
            # Reset for new analysis
            self.issues = []
            self.fixes_applied = []
            
            # Basic file validation
            if not os.path.exists(file_path):
                return self._create_error_result(f"File not found: {file_path}")
            
            # Open and analyze Word document
            doc = Document(file_path)
            
            doc_info = self._extract_document_info(doc)
            text_content = self._extract_text_content(doc)
            images = self._extract_images(doc)
            structure = self._analyze_structure(doc)
            styles = self._analyze_styles(doc)
            tables = self._analyze_tables(doc)
            
            # Perform accessibility checks
            self._check_document_structure(doc_info, structure)
            self._check_text_accessibility(text_content, styles)
            self._check_image_accessibility(images)
            self._check_table_accessibility(tables)
            self._check_heading_structure(structure)
            self._check_style_accessibility(styles)
            
            # Calculate accessibility score
            score = self._calculate_accessibility_score()
            
            # Apply fixes if requested
            if apply_fixes:
                output_path = self._apply_automatic_fixes(file_path, doc)
            else:
                output_path = None
            
            return self._create_analysis_result(
                file_path=file_path,
                output_path=output_path,
                score=score,
                doc_info=doc_info,
                text_content=text_content,
                images=images,
                structure=structure,
                styles=styles,
                tables=tables
            )
            
        except Exception as e:
            logger.error(f"Error analyzing Word document: {str(e)}")
            return self._create_error_result(f"Analysis failed: {str(e)}")
    
    def _extract_document_info(self, doc: Document) -> Dict[str, Any]:
        """Extract basic document metadata and properties"""
        info = {
            "title": None,
            "author": None,
            "subject": None,
            "paragraphs": len(doc.paragraphs),
            "sections": len(doc.sections),
            "has_toc": False,
            "language": None
        }
        
        try:
            # Core document properties
            if doc.core_properties:
                info["title"] = doc.core_properties.title
                info["author"] = doc.core_properties.author
                info["subject"] = doc.core_properties.subject
                info["language"] = doc.core_properties.language
            
            # Check for table of contents (simplified check)
            for para in doc.paragraphs:
                if para.style.name.startswith('TOC'):
                    info["has_toc"] = True
                    break
                    
        except Exception as e:
            logger.warning(f"Error extracting document info: {str(e)}")
        
        return info
    
    def _extract_text_content(self, doc: Document) -> Dict[str, Any]:
        """Extract and analyze text content"""
        content = {
            "paragraphs": [],
            "total_chars": 0,
            "word_count": 0,
            "reading_level": None,
            "all_caps_count": 0,
            "small_text_count": 0
        }
        
        try:
            full_text = ""
            
            for para_idx, para in enumerate(doc.paragraphs):
                para_text = para.text.strip()
                if not para_text:
                    continue
                
                para_info = {
                    "index": para_idx,
                    "text": para_text,
                    "style": para.style.name if para.style else "Normal",
                    "char_count": len(para_text),
                    "is_all_caps": para_text.isupper() and len(para_text) > 5,
                    "runs": []
                }
                
                # Analyze runs for formatting
                for run in para.runs:
                    if run.text.strip():
                        run_info = {
                            "text": run.text,
                            "font_name": run.font.name,
                            "font_size": run.font.size.pt if run.font.size else None,
                            "bold": run.font.bold,
                            "italic": run.font.italic,
                            "underline": run.font.underline,
                            "color": self._get_color_info(run.font.color)
                        }
                        para_info["runs"].append(run_info)
                        
                        # Check for small text
                        if run_info["font_size"] and run_info["font_size"] < 10:
                            content["small_text_count"] += 1
                
                # Check for all caps
                if para_info["is_all_caps"]:
                    content["all_caps_count"] += 1
                
                content["paragraphs"].append(para_info)
                full_text += para_text + " "
                content["total_chars"] += len(para_text)
            
            # Calculate reading level if textstat is available
            if textstat and len(full_text) > 100:
                content["reading_level"] = {
                    "flesch_reading_ease": textstat.flesch_reading_ease(full_text),
                    "flesch_kincaid_grade": textstat.flesch_kincaid_grade(full_text),
                    "automated_readability_index": textstat.automated_readability_index(full_text)
                }
                content["word_count"] = len(full_text.split())
            
        except Exception as e:
            logger.warning(f"Error extracting text content: {str(e)}")
        
        return content
    
    def _extract_images(self, doc: Document) -> List[Dict[str, Any]]:
        """Extract and analyze images from document"""
        images = []
        
        try:
            # Get all relationships to find images
            for rel in doc.part.rels.values():
                if "image" in rel.target_ref:
                    image_info = {
                        "filename": rel.target_ref.split('/')[-1],
                        "has_alt_text": False,
                        "alt_text": "",
                        "is_decorative": False
                    }
                    images.append(image_info)
            
            # Check inline shapes for alt text
            for para in doc.paragraphs:
                for run in para.runs:
                    if run._element.xpath('.//pic:pic'):
                        # This is a complex XPath query to find images with alt text
                        # For simplicity, we'll mark as needing manual review
                        pass
            
        except Exception as e:
            logger.warning(f"Error extracting images: {str(e)}")
        
        return images
    
    def _analyze_structure(self, doc: Document) -> Dict[str, Any]:
        """Analyze document structure"""
        structure = {
            "headings": [],
            "heading_levels": [],
            "has_proper_hierarchy": True,
            "lists": [],
            "links": []
        }
        
        try:
            previous_level = 0
            
            for para in doc.paragraphs:
                style_name = para.style.name if para.style else "Normal"
                
                # Check for headings
                if style_name.startswith('Heading'):
                    try:
                        level = int(style_name.split()[-1])
                        structure["headings"].append({
                            "text": para.text,
                            "level": level,
                            "style": style_name
                        })
                        structure["heading_levels"].append(level)
                        
                        # Check hierarchy (shouldn't skip levels)
                        if previous_level > 0 and level > previous_level + 1:
                            structure["has_proper_hierarchy"] = False
                        previous_level = level
                        
                    except (ValueError, IndexError):
                        pass
                
                # Check for lists (basic detection)
                if style_name.startswith('List') or para.text.strip().startswith(('•', '-', '*')):
                    structure["lists"].append({
                        "text": para.text[:50] + "..." if len(para.text) > 50 else para.text,
                        "style": style_name
                    })
                
                # Check for hyperlinks
                for run in para.runs:
                    if hasattr(run, 'hyperlink') and run.hyperlink:
                        structure["links"].append({
                            "text": run.text,
                            "address": getattr(run.hyperlink, 'address', 'Unknown')
                        })
            
        except Exception as e:
            logger.warning(f"Error analyzing structure: {str(e)}")
        
        return structure
    
    def _analyze_styles(self, doc: Document) -> Dict[str, Any]:
        """Analyze document styles"""
        styles_info = {
            "available_styles": [],
            "used_styles": set(),
            "custom_styles": [],
            "heading_styles": []
        }
        
        try:
            # Get all available styles
            for style in doc.styles:
                style_info = {
                    "name": style.name,
                    "type": str(style.type) if hasattr(style, 'type') else "Unknown",
                    "builtin": style.builtin if hasattr(style, 'builtin') else False
                }
                styles_info["available_styles"].append(style_info)
                
                if style.name.startswith('Heading'):
                    styles_info["heading_styles"].append(style.name)
                
                if hasattr(style, 'builtin') and not style.builtin:
                    styles_info["custom_styles"].append(style.name)
            
            # Track used styles
            for para in doc.paragraphs:
                if para.style:
                    styles_info["used_styles"].add(para.style.name)
            
            styles_info["used_styles"] = list(styles_info["used_styles"])
            
        except Exception as e:
            logger.warning(f"Error analyzing styles: {str(e)}")
        
        return styles_info
    
    def _analyze_tables(self, doc: Document) -> List[Dict[str, Any]]:
        """Analyze tables in the document"""
        tables_info = []
        
        try:
            for table_idx, table in enumerate(doc.tables):
                table_info = {
                    "index": table_idx,
                    "rows": len(table.rows),
                    "cols": len(table.columns) if table.rows else 0,
                    "has_header_row": False,
                    "cells": []
                }
                
                # Analyze table structure
                if table.rows:
                    # Check if first row looks like headers
                    first_row = table.rows[0]
                    header_indicators = 0
                    
                    for cell in first_row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            # Simple heuristic for header detection
                            for para in cell.paragraphs:
                                for run in para.runs:
                                    if run.font.bold:
                                        header_indicators += 1
                                        break
                    
                    table_info["has_header_row"] = header_indicators > 0
                
                tables_info.append(table_info)
                
        except Exception as e:
            logger.warning(f"Error analyzing tables: {str(e)}")
        
        return tables_info
    
    def _get_color_info(self, color) -> Dict[str, Any]:
        """Extract color information from a run"""
        color_info = {
            "type": "auto",
            "rgb": None,
            "theme_color": None
        }
        
        try:
            if color and hasattr(color, 'rgb') and color.rgb:
                color_info["type"] = "rgb"
                color_info["rgb"] = str(color.rgb)
            elif color and hasattr(color, 'theme_color') and color.theme_color:
                color_info["type"] = "theme"
                color_info["theme_color"] = str(color.theme_color)
        except Exception:
            pass
        
        return color_info
    
    def _check_document_structure(self, doc_info: Dict, structure: Dict):
        """Check document structure accessibility"""
        
        # Check for document title
        if not doc_info.get("title"):
            self.issues.append({
                "type": "structure",
                "severity": "high",
                "issue": "Missing document title",
                "description": "Document should have a descriptive title in properties",
                "recommendation": "Add a descriptive title in File > Info > Properties",
                "wcag_criterion": "2.4.2 Page Titled"
            })
        
        # Check for document language
        if not doc_info.get("language"):
            self.issues.append({
                "type": "structure",
                "severity": "high", 
                "issue": "Missing document language",
                "description": "Document should specify its primary language",
                "recommendation": "Set document language in Review > Language",
                "wcag_criterion": "3.1.1 Language of Page"
            })
        
        # Check heading structure
        if not structure.get("headings") and doc_info.get("paragraphs", 0) > 10:
            self.issues.append({
                "type": "structure",
                "severity": "medium",
                "issue": "Missing heading structure",
                "description": "Long documents should use headings to organize content",
                "recommendation": "Use Heading 1, Heading 2, etc. styles to structure content",
                "wcag_criterion": "1.3.1 Info and Relationships"
            })
        
        # Check heading hierarchy
        if structure.get("headings") and not structure.get("has_proper_hierarchy"):
            self.issues.append({
                "type": "structure", 
                "severity": "medium",
                "issue": "Improper heading hierarchy",
                "description": "Heading levels should not skip (e.g., H1 directly to H3)",
                "recommendation": "Use sequential heading levels (H1, H2, H3, etc.)",
                "wcag_criterion": "1.3.1 Info and Relationships"
            })
    
    def _check_text_accessibility(self, text_content: Dict, styles: Dict):
        """Check text accessibility issues"""
        
        # Check for excessive all caps text
        if text_content.get("all_caps_count", 0) > 0:
            self.issues.append({
                "type": "text",
                "severity": "medium",
                "issue": f"{text_content['all_caps_count']} paragraphs in all caps",
                "description": "All caps text is harder to read and may be interpreted as shouting",
                "recommendation": "Use sentence case with bold or emphasis for importance",
                "wcag_criterion": "1.4.8 Visual Presentation"
            })
        
        # Check for small text
        if text_content.get("small_text_count", 0) > 0:
            self.issues.append({
                "type": "text",
                "severity": "medium",
                "issue": f"{text_content['small_text_count']} instances of small text (< 10pt)",
                "description": "Very small text may be difficult for some users to read",
                "recommendation": "Use minimum 10pt font size, preferably 12pt or larger",
                "wcag_criterion": "1.4.12 Text Spacing"
            })
        
        # Check reading level
        reading_level = text_content.get("reading_level")
        if reading_level:
            grade_level = reading_level.get("flesch_kincaid_grade", 0)
            if grade_level > 12:
                self.issues.append({
                    "type": "text",
                    "severity": "low",
                    "issue": f"Complex reading level: Grade {grade_level:.1f}",
                    "description": "Content may be difficult for some users to understand",
                    "recommendation": "Consider simplifying language where appropriate",
                    "wcag_criterion": "3.1.5 Reading Level"
                })
    
    def _check_image_accessibility(self, images: List[Dict]):
        """Check image accessibility"""
        
        for img in images:
            if not img.get("has_alt_text") and not img.get("is_decorative"):
                self.issues.append({
                    "type": "images",
                    "severity": "high",
                    "issue": f"Image '{img['filename']}' missing alternative text",
                    "description": "Images need alternative text for screen readers",
                    "recommendation": "Right-click image > Format Picture > Alt Text and add description",
                    "wcag_criterion": "1.1.1 Non-text Content"
                })
    
    def _check_table_accessibility(self, tables: List[Dict]):
        """Check table accessibility"""
        
        for table in tables:
            if not table.get("has_header_row") and table.get("rows", 0) > 1:
                self.issues.append({
                    "type": "tables",
                    "severity": "medium",
                    "issue": f"Table {table['index'] + 1} missing header row",
                    "description": "Tables should have header rows to identify columns",
                    "recommendation": "Select first row and use Table Design > Header Row",
                    "wcag_criterion": "1.3.1 Info and Relationships"
                })
    
    def _check_heading_structure(self, structure: Dict):
        """Check heading structure specifically"""
        headings = structure.get("headings", [])
        
        if headings:
            # Check if document starts with H1
            if headings[0]["level"] != 1:
                self.issues.append({
                    "type": "structure",
                    "severity": "medium",
                    "issue": "Document doesn't start with Heading 1",
                    "description": "Documents should typically start with a Heading 1",
                    "recommendation": "Use Heading 1 for the main document title",
                    "wcag_criterion": "1.3.1 Info and Relationships"
                })
    
    def _check_style_accessibility(self, styles: Dict):
        """Check style usage for accessibility"""
        
        used_styles = styles.get("used_styles", [])
        
        # Check if only Normal style is used (poor structure)
        if len(used_styles) == 1 and "Normal" in used_styles:
            self.issues.append({
                "type": "structure",
                "severity": "medium",
                "issue": "Only Normal style used",
                "description": "Using only Normal style makes content structure unclear",
                "recommendation": "Use Heading styles and other built-in styles for structure",
                "wcag_criterion": "1.3.1 Info and Relationships"
            })
    
    def _calculate_accessibility_score(self) -> int:
        """Calculate overall accessibility score"""
        if not self.issues:
            return 100
        
        severity_weights = {
            "critical": 20,
            "high": 10,
            "medium": 5,
            "low": 2
        }
        
        total_deductions = sum(
            severity_weights.get(issue["severity"], 5) 
            for issue in self.issues
        )
        
        score = max(0, 100 - total_deductions)
        return score
    
    def _apply_automatic_fixes(self, file_path: str, doc: Document) -> str:
        """Apply automatic fixes to the document"""
        
        try:
            # Create output path
            input_path = Path(file_path)
            output_path = input_path.parent / "output" / f"accessible_{input_path.name}"
            output_path.parent.mkdir(exist_ok=True)
            
            # Apply fixes
            fixes_count = 0
            
            # Fix 1: Add document title if missing
            if not doc.core_properties.title:
                # Use first heading or first paragraph as title
                title = "Untitled Document"
                for para in doc.paragraphs:
                    if para.text.strip():
                        title = para.text.strip()[:100]  # First 100 chars
                        break
                doc.core_properties.title = title
                fixes_count += 1
                self.fixes_applied.append({
                    "type": "metadata",
                    "description": f"Added document title: {title}"
                })
            
            # Fix 2: Convert all caps paragraphs to sentence case
            for para in doc.paragraphs:
                if para.text.isupper() and len(para.text) > 5:
                    # Convert to title case
                    para.text = para.text.title()
                    fixes_count += 1
                    self.fixes_applied.append({
                        "type": "text",
                        "description": f"Converted all caps text to title case"
                    })
            
            # Fix 3: Increase very small font sizes
            for para in doc.paragraphs:
                for run in para.runs:
                    if run.font.size and run.font.size.pt < 10:
                        run.font.size = Pt(10)
                        fixes_count += 1
                        self.fixes_applied.append({
                            "type": "text",
                            "description": f"Increased font size to 10pt minimum"
                        })
            
            # Save the modified document
            doc.save(str(output_path))
            
            logger.info(f"Applied {fixes_count} automatic fixes to {output_path}")
            return str(output_path)
            
        except Exception as e:
            logger.error(f"Error applying fixes: {str(e)}")
            self.fixes_applied.append({
                "type": "error",
                "description": f"Could not apply automatic fixes: {str(e)}"
            })
            return None
    
    def _create_analysis_result(self, file_path: str, score: int, output_path: str = None, **kwargs) -> Dict[str, Any]:
        """Create standardized analysis result"""
        return {
            "success": True,
            "file_path": file_path,
            "output_path": output_path,
            "file_type": "docx",
            "accessibility_score": score,
            "total_issues": len(self.issues),
            "critical_issues": len([i for i in self.issues if i["severity"] == "critical"]),
            "high_issues": len([i for i in self.issues if i["severity"] == "high"]),
            "medium_issues": len([i for i in self.issues if i["severity"] == "medium"]),
            "low_issues": len([i for i in self.issues if i["severity"] == "low"]),
            "issues": self.issues,
            "fixes_applied": self.fixes_applied,
            "document_info": kwargs.get("doc_info", {}),
            "recommendations": self._generate_recommendations(),
            "wcag_compliance": self._assess_wcag_compliance()
        }
    
    def _create_error_result(self, error_message: str) -> Dict[str, Any]:
        """Create error result"""
        return {
            "success": False,
            "error": error_message,
            "file_type": "docx",
            "accessibility_score": 0,
            "total_issues": 0,
            "issues": [],
            "fixes_applied": []
        }
    
    def _generate_recommendations(self) -> List[str]:
        """Generate prioritized recommendations"""
        recommendations = []
        
        critical_issues = [i for i in self.issues if i["severity"] == "critical"]
        high_issues = [i for i in self.issues if i["severity"] == "high"]
        
        if critical_issues:
            recommendations.append("❗ CRITICAL: Address document structure and accessibility features")
            
        if high_issues:
            recommendations.append("🔥 HIGH PRIORITY: Add document metadata and alt text for images")
            
        recommendations.extend([
            "📝 Use built-in styles (Heading 1, Heading 2, etc.) for structure",
            "🖼️ Add alternative text to all informative images",
            "📊 Use table headers for data tables",
            "✅ Run Word's built-in Accessibility Checker"
        ])
        
        return recommendations
    
    def _assess_wcag_compliance(self) -> Dict[str, str]:
        """Assess WCAG 2.1 Level AA compliance"""
        critical_issues = [i for i in self.issues if i["severity"] == "critical"]
        high_issues = [i for i in self.issues if i["severity"] == "high"]
        
        if critical_issues:
            return {
                "level": "Non-compliant",
                "status": "Major accessibility barriers present",
                "next_steps": "Address critical issues before using in courses"
            }
        elif high_issues:
            return {
                "level": "Partially compliant",
                "status": "Some accessibility improvements needed", 
                "next_steps": "Address high-priority issues for better compliance"
            }
        else:
            return {
                "level": "Mostly compliant",
                "status": "Meets most WCAG 2.1 AA requirements",
                "next_steps": "Address remaining minor issues for full compliance"
            }